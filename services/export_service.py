"""Export optimized resume to PDF and DOCX."""

import io
import os
import tempfile
from typing import Any


class ExportService:
    def export_docx(self, resume_doc: dict[str, Any], filename: str = "optimized_resume.docx") -> str:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        contact = resume_doc.get("contact", {})

        name = doc.add_heading(contact.get("name", "Resume"), 0)
        name.alignment = 1

        line_parts = [
            contact.get("email"),
            contact.get("phone"),
            contact.get("linkedin"),
            contact.get("location"),
        ]
        contact_line = " | ".join(p for p in line_parts if p)
        if contact_line:
            p = doc.add_paragraph(contact_line)
            p.alignment = 1

        if resume_doc.get("summary"):
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(resume_doc["summary"])

        skills = resume_doc.get("skills", {})
        if skills:
            doc.add_heading("Skills", level=1)
            if isinstance(skills, dict):
                for cat, items in skills.items():
                    doc.add_paragraph(f"{cat.replace('_', ' ').title()}: {', '.join(items)}")
            else:
                doc.add_paragraph(", ".join(skills))

        if resume_doc.get("experience"):
            doc.add_heading("Experience", level=1)
            for exp in resume_doc["experience"]:
                hdr = exp.get("title", "")
                if exp.get("company"):
                    hdr += f" — {exp['company']}"
                if exp.get("dates"):
                    hdr += f" ({exp['dates']})"
                doc.add_paragraph(hdr).runs[0].bold = True
                for b in exp.get("bullets", []):
                    doc.add_paragraph(b, style="List Bullet")

        if resume_doc.get("education"):
            doc.add_heading("Education", level=1)
            for ed in resume_doc["education"]:
                line = f"{ed.get('degree', '')}, {ed.get('institution', '')}"
                if ed.get("years"):
                    line += f" ({ed['years']})"
                doc.add_paragraph(line)

        if resume_doc.get("projects"):
            doc.add_heading("Projects", level=1)
            for pr in resume_doc["projects"]:
                doc.add_paragraph(pr.get("title", "Project")).runs[0].bold = True
                for b in pr.get("bullets", []):
                    doc.add_paragraph(b, style="List Bullet")

        if resume_doc.get("certifications"):
            doc.add_heading("Certifications", level=1)
            for c in resume_doc["certifications"]:
                doc.add_paragraph(c, style="List Bullet")

        path = os.path.join(tempfile.gettempdir(), filename)
        doc.save(path)
        return path

    def export_pdf(self, resume_doc: dict[str, Any], filename: str = "optimized_resume.pdf") -> str:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem

        path = os.path.join(tempfile.gettempdir(), filename)
        doc = SimpleDocTemplate(path, pagesize=letter, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "Title", parent=styles["Heading1"], fontSize=16, spaceAfter=6,
        )
        h_style = ParagraphStyle(
            "H2", parent=styles["Heading2"], fontSize=12, spaceBefore=12, spaceAfter=6,
        )
        body = styles["BodyText"]
        story = []

        contact = resume_doc.get("contact", {})
        story.append(Paragraph(self._xml(contact.get("name", "Resume")), title_style))
        parts = [contact.get("email"), contact.get("phone"), contact.get("linkedin")]
        line = " | ".join(p for p in parts if p)
        if line:
            story.append(Paragraph(self._xml(line), body))
        story.append(Spacer(1, 12))

        if resume_doc.get("summary"):
            story.append(Paragraph("PROFESSIONAL SUMMARY", h_style))
            story.append(Paragraph(self._xml(resume_doc["summary"]), body))

        skills = resume_doc.get("skills", {})
        if skills:
            story.append(Paragraph("SKILLS", h_style))
            if isinstance(skills, dict):
                for cat, items in skills.items():
                    story.append(Paragraph(
                        self._xml(f"{cat.replace('_', ' ').title()}: {', '.join(items)}"),
                        body,
                    ))
            else:
                story.append(Paragraph(self._xml(", ".join(skills)), body))

        for section, key in [("EXPERIENCE", "experience"), ("EDUCATION", "education"), ("PROJECTS", "projects")]:
            items = resume_doc.get(key, [])
            if not items:
                continue
            story.append(Paragraph(section, h_style))
            for item in items:
                if key == "experience":
                    hdr = item.get("title", "")
                    if item.get("company"):
                        hdr += f" — {item['company']}"
                    if item.get("dates"):
                        hdr += f" ({item['dates']})"
                    story.append(Paragraph(self._xml(hdr), body))
                    bullets = [ListItem(Paragraph(self._xml(b), body)) for b in item.get("bullets", [])]
                    if bullets:
                        story.append(ListFlowable(bullets, bulletType="bullet"))
                elif key == "education":
                    line = f"{item.get('degree', '')}, {item.get('institution', '')}"
                    if item.get("years"):
                        line += f" ({item['years']})"
                    story.append(Paragraph(self._xml(line), body))
                else:
                    story.append(Paragraph(self._xml(item.get("title", "")), body))
                    for b in item.get("bullets", []):
                        story.append(Paragraph(self._xml(f"• {b}"), body))

        doc.build(story)
        return path

    def _xml(self, text: str) -> str:
        if not text:
            return ""
        return (
            str(text)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
